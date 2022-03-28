from datetime import date

from django.http import HttpResponse, HttpResponseRedirect
from django.shortcuts import render, redirect
from django.urls import reverse
from django.views import View

from .models import usercanvas, questions, selectiveQuestion, usercanvasSelective, userInformation, userProfile, \
    docQuestions, english_docQuestions, class_suggesstion

from django.contrib.auth.decorators import login_required

from django.core.mail import send_mail
from .permission import allowed_users
from . import forms

from bijoytounicode import bijoy2unicode


@login_required(login_url="/accounts/login/")
# @allowed_users(allowed_roles=['staff'])
def homepage(request):
    try:
        user = userInformation.objects.get(user=request.user)
    except:
        return redirect('articles:userInfo')
    return render(request, 'main/homepage.html', {'user': user})


@login_required(login_url="/account/login/")
def userInfo(request):
    try:
        user = userInformation.objects.get(user=request.user)
        if user is not None:
            return redirect('articles:list')
    except:
        print("Stay here")
    form = forms.userInformation()
    if request.method == 'POST':
        form = forms.userInformation(request.POST, request.FILES)
        if form.is_valid():
            instance = form.save(commit=False)
            instance.user = request.user
            instance.save()
            url = reverse('articles:list')
            next = request.POST.get('next', '/')
            return HttpResponseRedirect(url)
    return render(request, 'main/updateInformation.html', {'form': form})


@login_required(login_url="/account/login/")
def contact(request):
    if request.method == "POST":
        message_name = request.POST['message-name']
        message_email = request.POST['message-email']
        message_phone = request.POST['message-phone']
        message = request.POST['message-message']

        # send email
        send_mail(
            'Mail Sent By ' + message_name,  # subject
            '\n' + 'Senders Phone: ' + message_phone + '\nSenders Email: ' + message_email + ' \nMessage: ' + message,
            # message
            message_email,  # from mail
            ['rashikbuksh71@gmail.com'],  # tomail
        )
        return render(request, 'main/contact.html', {'message_name': message_name})
    else:
        return render(request, 'main/contact.html')


@login_required(login_url="/account/login/")
def classAndSubjects(request):
    classAndSubjects.classInput = ""
    classAndSubjects.subjectInput = ""
    classAndSubjects.classInput = request.POST.get('class-Input')
    classAndSubjects.subjectInput = request.POST.get('subject-Input')
    if classAndSubjects.classInput and classAndSubjects.subjectInput is not None:
        return redirect('articles:questions')
    return render(request, 'main/classAndSubjects.html')


@login_required(login_url="/account/login/")
def Canvas(request):
    canvass = usercanvas.objects.filter(user=request.user)
    canvass1 = usercanvasSelective.objects.filter(user=request.user)
    context = {'canvass': canvass, 'canvass1': canvass1}
    return render(request, "main/canvas.html", context)


@login_required(login_url="/account/login/")
def removequestion(request, pk):
    instance = usercanvas.objects.get(id=pk)
    instance.delete()
    return redirect('articles:canvas')


@login_required(login_url="/account/login/")
def removequestionSelective(request, pk):
    instance = usercanvasSelective.objects.get(id=pk)
    instance.delete()
    return redirect('articles:canvas')


@login_required(login_url="/account/login/")
def questionsss(request):
    questionss = questions.objects.filter(classes=int(classAndSubjects.classInput),
                                          subject=classAndSubjects.subjectInput)
    questionsSelective = selectiveQuestion.objects.filter(classes=int(classAndSubjects.classInput),
                                                          subject=classAndSubjects.subjectInput)
    context = {'questionss': questionss, 'questionsSelective': questionsSelective}
    return render(request, 'main/questionAdd.html', context)


@login_required(login_url="/account/login/")
def addquestion(request, pk):
    instance = questions.objects.get(id=pk)
    instance2 = usercanvas.objects.create(user=request.user, scenario=instance.scenario, ques_img=instance.ques_img,
                                          q_a=instance.q_a, q_b=instance.q_b, q_c=instance.q_c, q_d=instance.q_d)
    instance2.save()
    return redirect('articles:canvas')


@login_required(login_url="/account/login/")
def addquestionSelective(request, pk):
    instance = selectiveQuestion.objects.get(id=pk)
    instance2 = usercanvasSelective.objects.create(user=request.user, scenario=instance.scenario,
                                                   ques_img=instance.ques_img,
                                                   q_a=instance.q_a, q_b=instance.q_b, q_c=instance.q_c,
                                                   q_d=instance.q_d)
    instance2.save()
    return redirect('articles:canvas')


@login_required(login_url="/account/login/")
def editquestion(request, pk):
    instance = usercanvas.objects.get(id=pk)
    context = {}
    context['editCanvas'] = instance

    scenario = request.POST.get('canv-scenario')
    qa = request.POST.get('canv-qa')
    qb = request.POST.get('canv-qb')
    qc = request.POST.get('canv-qc')
    qd = request.POST.get('canv-qd')

    if scenario is not None and scenario != '':
        instance.scenario = scenario

    if qa is not None and qa != '':
        instance.q_a = qa

    if qb is not None and qb != '':
        instance.q_b = qb

    if qc is not None and qc != '':
        instance.q_c = qc

    if qd is not None and qd != '':
        instance.q_d = qd

    instance.save()
    return render(request, 'main/editQuestion.html', context)


@login_required(login_url="/account/login/")
def editquestionSelective(request, pk):
    instance = usercanvasSelective.objects.get(id=pk)
    context = {}
    context['editCanvas'] = instance

    scenario = request.POST.get('canv-scenario')
    qa = request.POST.get('canv-qa')
    qb = request.POST.get('canv-qb')
    qc = request.POST.get('canv-qc')
    qd = request.POST.get('canv-qd')

    if scenario is not None and scenario != '':
        instance.scenario = scenario

    if qa is not None and qa != '':
        instance.q_a = qa

    if qb is not None and qb != '':
        instance.q_b = qb

    if qc is not None and qc != '':
        instance.q_c = qc

    if qd is not None and qd != '':
        instance.q_d = qd

    instance.save()
    return render(request, 'main/editQuestion.html', context)


@login_required(login_url="/account/login/")
def questionsGenerate(request):
    userInfo = userInformation.objects.get(user=request.user)
    count = 0
    canvass = usercanvas.objects.filter(user=request.user)
    for i in canvass:
        count += 1
    canvassSelective = usercanvasSelective.objects.filter(user=request.user)
    for i in canvassSelective:
        count += 1
    if userInfo.point >= count:
        userInfo.point = userInfo.point - count
        userInfo.save()
    else:
        return redirect('articles:list')
    context = {'canvass': canvass, 'canvassSelective': canvassSelective}
    return render(request, 'main/QuestionGenerate.html', context)


@login_required(login_url="/account/login/")
@allowed_users(allowed_roles=['staff'])
def questioninput(request):
    return render(request, 'main/QuestionInput.html')


@login_required(login_url="/account/login/")
@allowed_users(allowed_roles=['staff'])
def bijoyinput(request):
    scenario = request.POST.get('scenario')
    qa = request.POST.get('qa')
    qb = request.POST.get('qb')
    qc = request.POST.get('qc')
    qd = request.POST.get('qd')
    if scenario is not None:
        convertedScenario = bijoy2unicode(scenario)
    if qa is not None:
        convertedqa = bijoy2unicode(qa)
    if qb is not None:
        convertedqb = bijoy2unicode(qb)
    if qc is not None:
        convertedqc = bijoy2unicode(qc)
    if qd is not None:
        convertedqd = bijoy2unicode(qd)
    form = forms.questionInputBijoy()
    if request.method == 'POST':
        form = forms.questionInputBijoy(request.POST, request.FILES)
        if form.is_valid():
            instance = form.save(commit=False)
            instance.scenario = convertedScenario
            instance.q_a = convertedqa
            instance.q_b = convertedqb
            instance.q_c = convertedqc
            instance.q_d = convertedqd
            instance.save()
            url = reverse('articles:questioninput')
            next = request.POST.get('next', '/')
            return HttpResponseRedirect(url)
    return render(request, 'main/bijoyquestions.html', {'form': form})


@login_required(login_url="/account/login/")
@allowed_users(allowed_roles=['staff'])
def unicodeinput(request):
    form = forms.questionInput()
    if request.method == 'POST':
        form = forms.questionInput(request.POST, request.FILES)
        if form.is_valid():
            instance = form.save(commit=False)
            instance.save()
            url = reverse('articles:questioninput')
            next = request.POST.get('next', '/')
            return HttpResponseRedirect(url)
    return render(request, 'main/unicodequestions.html', {'form': form})


@login_required(login_url="/account/login/")
@allowed_users(allowed_roles=['staff'])
def bijoyinputSelective(request):
    scenario = request.POST.get('scenario')
    qa = request.POST.get('qa')
    qb = request.POST.get('qb')
    qc = request.POST.get('qc')
    qd = request.POST.get('qd')
    if scenario is not None:
        convertedScenario = bijoy2unicode(scenario)
    if qa is not None:
        convertedqa = bijoy2unicode(qa)
    if qb is not None:
        convertedqb = bijoy2unicode(qb)
    if qc is not None:
        convertedqc = bijoy2unicode(qc)
    if qd is not None:
        convertedqd = bijoy2unicode(qd)
    form = forms.questionInputBijoySelective()
    if request.method == 'POST':
        form = forms.questionInputBijoySelective(request.POST, request.FILES)
        if form.is_valid():
            instance = form.save(commit=False)
            instance.scenario = convertedScenario
            instance.q_a = convertedqa
            instance.q_b = convertedqb
            instance.q_c = convertedqc
            instance.q_d = convertedqd
            instance.save()
            url = reverse('articles:questioninput')
            next = request.POST.get('next', '/')
            return HttpResponseRedirect(url)
    return render(request, 'main/bijoySelective.html', {'form': form})


@login_required(login_url="/account/login/")
@allowed_users(allowed_roles=['staff'])
def unicodeinputSelective(request):
    form = forms.questionInputSelective()
    if request.method == 'POST':
        form = forms.questionInputSelective(request.POST, request.FILES)
        if form.is_valid():
            instance = form.save(commit=False)
            instance.save()
            url = reverse('articles:questioninput')
            next = request.POST.get('next', '/')
            return HttpResponseRedirect(url)
    return render(request, 'main/unicodeSelective.html', {'form': form})


@login_required(login_url="/account/login/")
def purchasePoint(request):
    form = forms.paymentInformation()
    if request.method == 'POST':
        form = forms.paymentInformation(request.POST, request.FILES)
        if form.is_valid():
            instance = form.save(commit=False)
            instance.user = request.user
            instance.save()
            url = reverse('articles:list')
            next = request.POST.get('next', '/')
            return HttpResponseRedirect(url)
    return render(request, 'main/PurchasePoint.html', {'form': form})


@login_required(login_url="/account/login/")
def updateUserProfile(request, pk):
    instance = userProfile.objects.get(id=pk)
    instance.save()
    updateUserProfile.primaryKey = pk
    return redirect('articles:user_profile')


@login_required(login_url="/account/login/")
def UserProfile(request):
    primaryKey = request.POST.get('primaryKey')
    if primaryKey is None:
        primaryKey = 1
    instance = userProfile.objects.get(id=primaryKey)
    UserProfile.user_name = request.POST.get('user_name')
    UserProfile.phone = request.POST.get('phone_number')
    UserProfile.address = request.POST.get('address')
    UserProfile.bio = request.POST.get('bio')
    UserProfile.image = request.FILES.get('image')

    if UserProfile.user_name is not None and UserProfile.user_name != '':
        instance.user_name = UserProfile.user_name

    if UserProfile.phone is not None and UserProfile.phone != '':
        instance.user_phone = UserProfile.phone

    if UserProfile.address is not None and UserProfile.address != '':
        instance.user_address = UserProfile.address

    if UserProfile.bio is not None and UserProfile.bio != '':
        instance.bio = UserProfile.bio

    if UserProfile.image is not None:
        instance.user_image = UserProfile.image
    instance.save()

    profile = userProfile.objects.all()
    profile2 = userProfile.objects.filter(user=request.user).first()
    b = ''
    if profile2 is None:
        b = 'NoData'

    context = {}
    context['profile'] = profile
    context['booli'] = b
    context['name'] = UserProfile.user_name
    context['phone'] = UserProfile.phone
    context['address'] = UserProfile.address
    context['bio'] = UserProfile.bio
    context['image'] = UserProfile.image

    return render(request, 'main/user_profile.html', context)


@login_required(login_url="/account/login/")
def createProfile(request):
    form = forms.UserProfile()
    if request.method == 'POST':
        form = forms.UserProfile(request.POST, request.FILES)
        if form.is_valid():
            instance = form.save(commit=False)
            instance.user = request.user
            instance.save()
            next = request.POST.get('next', '/')
            return HttpResponseRedirect(next)
    else:
        form = forms.UserProfile()
    return render(request, 'main/createprofile.html', {'form': form})


import docx


@login_required(login_url="/account/login/")
def openDocx(request, pk):
    openDocx.docOpen = request.POST.get('doc-Open')
    if openDocx.docOpen is not None:
        return redirect('articles:docQuestionPage')
    docc = docQuestions.objects.get(id=pk)
    context = {}
    context['doc'] = docc
    doc = docx.Document(docc.docs)
    fullText = []
    for para in doc.paragraphs:
        convertedpara = bijoy2unicode(para.text)
        fullText.append(convertedpara)
    context = {'data': fullText}

    return render(request, 'main/openDocx.html', context)


@login_required(login_url="/account/login/")
def lessThanSix(request):
    lessThanSix.classInput = ""
    lessThanSix.subjectInput = ""
    lessThanSix.classInput = request.POST.get('class-Input')
    lessThanSix.subjectInput = request.POST.get('subject-Input')
    if lessThanSix.classInput and lessThanSix.subjectInput is not None:
        return redirect('articles:showDoc')
    return render(request, 'main/classAndSubjectlessThenSix.html')


@login_required(login_url="/account/login/")
def showDoc(request):
    qlist = docQuestions.objects.filter(classes=lessThanSix.classInput, subject=lessThanSix.subjectInput)
    context = {'qlist': qlist}
    return render(request, 'main/showDocQuestions.html', context)


@login_required(login_url="/account/login/")
def docQuestionPage(request):
    context = {'docOpen': openDocx.docOpen}
    return render(request, 'main/docQuestionPage.html', context)


@login_required(login_url="/account/login/")
def openEnglishDocx(request, pk):
    openEnglishDocx.docOpen = request.POST.get('doc-Open')
    if openEnglishDocx.docOpen is not None:
        return redirect('articles:docEnglishQuestionPage')
    docc = english_docQuestions.objects.get(id=pk)
    context = {}
    context['doc'] = docc
    doc = docx.Document(docc.docs)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    context = {'data': fullText}

    return render(request, 'main/openEnglishDocx.html', context)


@login_required(login_url="/account/login/")
def englishQuestions(request):
    englishQuestions.classInput = ""
    englishQuestions.paperInput = ""
    englishQuestions.classInput = request.POST.get('class-Input')
    englishQuestions.paperInput = request.POST.get('paper-Input')
    print(englishQuestions.paperInput)
    if englishQuestions.classInput and englishQuestions.paperInput is not None:
        return redirect('articles:showEnglishDoc')
    return render(request, 'main/englishQuestions.html')


@login_required(login_url="/account/login/")
def showEnglishDoc(request):
    qlist = english_docQuestions.objects.filter(classes=englishQuestions.classInput, paper=englishQuestions.paperInput)
    context = {'qlist': qlist}
    return render(request, 'main/showEnglishDocQuestions.html', context)


@login_required(login_url="/account/login/")
def docEnglishQuestionPage(request):
    context = {'docOpen': openEnglishDocx.docOpen}
    return render(request, 'main/docQuestionPage.html', context)


@login_required(login_url="/account/login/")
def sugesstion_class(request):
    today = date.today()
    print(today.weekday())
    month, day = today.month, today.day
    instance = class_suggesstion.objects.filter(classes=1)
    context = {}
    context['weekday'] = 'None'
    if today.weekday() == 0:
        context['weekday'] = 'Monday'
    elif today.weekday() == 1:
        context['weekday'] = 'Tuesday'
    elif today.weekday() == 2:
        context['weekday'] = 'Wednesday'
    elif today.weekday() == 3:
        context['weekday'] = 'Thursday'
    elif today.weekday() == 4:
        context['weekday'] = 'Friday'
    elif today.weekday() == 5:
        context['weekday'] = 'Saturday'
    elif today.weekday() == 6:
        context['weekday'] = 'Sunday'

    context['bangla'] = 'None'
    context['english'] = 'None'
    context['math'] = 'None'
    context['science'] = 'None'
    context['social'] = 'None'
    context['islamReligion'] = 'None'
    context['hinduReligion'] = 'None'
    context['boudhuReligion'] = 'None'
    context['cristantianReligion'] = 'None'
    context['artsAndCrafts'] = 'None'
    context['physicalEducation'] = 'None'
    context['music'] = 'None'
    context['comment'] = 'None'
    for i in instance:
        if i.date.month == month and i.date.day == day:
            context['bangla'] = i.bangla
            context['english'] = i.english
            context['math'] = i.math
            context['science'] = i.science
            context['social'] = i.social
            context['islamReligion'] = i.islamReligion
            context['hinduReligion'] = i.hinduReligion
            context['boudhuReligion'] = i.boudhuReligion
            context['cristantianReligion'] = i.cristantianReligion
            context['artsAndCrafts'] = i.artsAndCrafts
            context['physicalEducation'] = i.physicalEducation
            context['music'] = i.music
            context['comment'] = i.comment
    return render(request, 'main/suggesstion_class.html', context)


@login_required(login_url="/account/login/")
@allowed_users(allowed_roles=['staff'])
def DocumentQuestionsform(request):
    form = forms.DocumentQuestions()
    if request.method == 'POST':
        form = forms.DocumentQuestions(request.POST, request.FILES)
        if form.is_valid():
            instance = form.save(commit=False)
            instance.save()
            url = reverse('articles:DocumentQuestionsform')
            next = request.POST.get('next', '/')
            return HttpResponseRedirect(url)
    return render(request, 'main/DocumentQuestionsForm.html', {'form': form})


@login_required(login_url="/account/login/")
@allowed_users(allowed_roles=['staff'])
def EnglishDocQuestionsform(request):
    form = forms.EnglishDocQuestions()
    if request.method == 'POST':
        form = forms.EnglishDocQuestions(request.POST, request.FILES)
        if form.is_valid():
            instance = form.save(commit=False)
            instance.save()
            url = reverse('articles:EnglishDocQuestionsform')
            next = request.POST.get('next', '/')
            return HttpResponseRedirect(url)
    return render(request, 'main/EnglishDocQuestionsForm.html', {'form': form})


@login_required(login_url="/account/login/")
@allowed_users(allowed_roles=['staff'])
def ClassSuggesstionform(request):
    form = forms.ClassSuggesstion()
    if request.method == 'POST':
        form = forms.ClassSuggesstion(request.POST, request.FILES)
        if form.is_valid():
            instance = form.save(commit=False)
            instance.save()
            url = reverse('articles:ClassSuggesstionform')
            next = request.POST.get('next', '/')
            return HttpResponseRedirect(url)
    return render(request, 'main/ClassSuggesstionForm.html', {'form': form})


@login_required(login_url="/account/login/")
def admitCard(request):
    admitCard.instName = request.POST.get('inst-name')
    admitCard.exam = request.POST.get('exam-name')
    admitCard.stdName = request.POST.get('student-name')
    admitCard.stdFatherName = request.POST.get('father-name')
    admitCard.stdMotherName = request.POST.get('mother-name')
    admitCard.stdClass = request.POST.get('class-number')
    admitCard.stdRoll = request.POST.get('roll-number')
    admitCard.headMaster = request.POST.get('head-master')

    if admitCard.instName and admitCard.exam and admitCard.stdName and admitCard.stdFatherName and admitCard.stdMotherName and admitCard.stdClass and admitCard.stdRoll is not None:
        return redirect('articles:admitCardGen')
    return render(request, 'main/admitCard.html')


@login_required(login_url="/account/login/")
def admitCardGen(request):
    context = {'instName': admitCard.instName, 'exam': admitCard.exam, 'stdName': admitCard.stdName,
               'stdFatherName': admitCard.stdFatherName, 'stdMotherName': admitCard.stdMotherName,
               'stdClass': admitCard.stdClass, 'stdRoll': admitCard.stdRoll, 'headMaster': admitCard.headMaster}
    return render(request, 'main/admitCardGen.html', context)


@login_required(login_url="/account/login/")
def routine(request):
    routine.date = request.POST.get('date')
    routine.weekday = request.POST.get('weekday')
    routine.date2 = request.POST.get('date2')
    routine.weekday2 = request.POST.get('weekday2')
    routine.date3 = request.POST.get('date3')
    routine.weekday3 = request.POST.get('weekday3')
    routine.date4 = request.POST.get('date4')
    routine.weekday4 = request.POST.get('weekday4')
    routine.date5 = request.POST.get('date5')
    routine.weekday5 = request.POST.get('weekday5')
    routine.date6 = request.POST.get('date6')
    routine.weekday6 = request.POST.get('weekday6')
    routine.date7 = request.POST.get('date7')
    routine.weekday7 = request.POST.get('weekday7')
    routine.p1 = request.POST.get('p1')
    routine.p2 = request.POST.get('p2')
    routine.p3 = request.POST.get('p3')
    routine.p4 = request.POST.get('p4')
    routine.p5 = request.POST.get('p5')
    routine.p6 = request.POST.get('p6')
    routine.p7 = request.POST.get('p7')
    routine.n1 = request.POST.get('n1')
    routine.n2 = request.POST.get('n2')
    routine.n3 = request.POST.get('n3')
    routine.n4 = request.POST.get('n4')
    routine.n5 = request.POST.get('n5')
    routine.n6 = request.POST.get('n6')
    routine.n7 = request.POST.get('n7')
    routine.c11 = request.POST.get('c11')
    routine.c12 = request.POST.get('c12')
    routine.c13 = request.POST.get('c13')
    routine.c14 = request.POST.get('c14')
    routine.c15 = request.POST.get('c15')
    routine.c16 = request.POST.get('c16')
    routine.c17 = request.POST.get('c17')
    routine.c21 = request.POST.get('c21')
    routine.c22 = request.POST.get('c22')
    routine.c23 = request.POST.get('c23')
    routine.c24 = request.POST.get('c24')
    routine.c25 = request.POST.get('c25')
    routine.c26 = request.POST.get('c26')
    routine.c27 = request.POST.get('c27')
    routine.c31 = request.POST.get('c31')
    routine.c32 = request.POST.get('c32')
    routine.c33 = request.POST.get('c33')
    routine.c34 = request.POST.get('c34')
    routine.c35 = request.POST.get('c35')
    routine.c36 = request.POST.get('c36')
    routine.c37 = request.POST.get('c37')
    routine.c41 = request.POST.get('c41')
    routine.c42 = request.POST.get('c42')
    routine.c43 = request.POST.get('c43')
    routine.c44 = request.POST.get('c44')
    routine.c45 = request.POST.get('c45')
    routine.c46 = request.POST.get('c46')
    routine.c47 = request.POST.get('c47')
    routine.c51 = request.POST.get('c51')
    routine.c52 = request.POST.get('c52')
    routine.c53 = request.POST.get('c53')
    routine.c54 = request.POST.get('c54')
    routine.c55 = request.POST.get('c55')
    routine.c56 = request.POST.get('c56')
    routine.c57 = request.POST.get('c57')
    routine.c61 = request.POST.get('c61')
    routine.c62 = request.POST.get('c62')
    routine.c63 = request.POST.get('c63')
    routine.c64 = request.POST.get('c64')
    routine.c65 = request.POST.get('c65')
    routine.c66 = request.POST.get('c66')
    routine.c67 = request.POST.get('c67')
    routine.c71 = request.POST.get('c71')
    routine.c72 = request.POST.get('c72')
    routine.c73 = request.POST.get('c73')
    routine.c74 = request.POST.get('c74')
    routine.c75 = request.POST.get('c75')
    routine.c76 = request.POST.get('c76')
    routine.c77 = request.POST.get('c77')
    routine.c81 = request.POST.get('c81')
    routine.c82 = request.POST.get('c82')
    routine.c83 = request.POST.get('c83')
    routine.c84 = request.POST.get('c84')
    routine.c85 = request.POST.get('c85')
    routine.c86 = request.POST.get('c86')
    routine.c87 = request.POST.get('c87')
    routine.c91 = request.POST.get('c91')
    routine.c92 = request.POST.get('c92')
    routine.c93 = request.POST.get('c93')
    routine.c94 = request.POST.get('c94')
    routine.c95 = request.POST.get('c95')
    routine.c96 = request.POST.get('c96')
    routine.c97 = request.POST.get('c97')
    routine.c101 = request.POST.get('c101')
    routine.c102 = request.POST.get('c102')
    routine.c103 = request.POST.get('c103')
    routine.c104 = request.POST.get('c104')
    routine.c105 = request.POST.get('c105')
    routine.c106 = request.POST.get('c106')
    routine.c107 = request.POST.get('c107')
    routine.schoolName = request.POST.get('school-name')
    if routine.weekday and routine.date and routine.schoolName:
        return redirect('articles:routinePrint')
    return render(request, 'main/routine.html')


def routinePrint(request):
    context = {'date': routine.date, 'date2': routine.date2, 'date3': routine.date3,
               'date4': routine.date4, 'date5': routine.date5, 'date6': routine.date6,
               'date7': routine.date7, 'weekday': routine.weekday,
               'weekday2': routine.weekday2, 'weekday3': routine.weekday3, 'weekday4': routine.weekday4,
               'weekday5': routine.weekday5, 'weekday6': routine.weekday6, 'weekday7': routine.weekday7,
               'p1': routine.p1, 'p2': routine.p2, 'p3': routine.p3, 'p4': routine.p4, 'p5': routine.p5,
               'p6': routine.p6, 'p7': routine.p7, 'n1': routine.n1, 'n2': routine.n2, 'n3': routine.n3,
               'n4': routine.n4, 'n5': routine.n5, 'n6': routine.n6, 'n7': routine.n7, 'c11': routine.c11,
               'c12': routine.c12, 'c13': routine.c13, 'c14': routine.c14, 'c15': routine.c15, 'c16': routine.c16,
               'c17': routine.c17, 'c21': routine.c21, 'c22': routine.c22, 'c23': routine.c23, 'c24': routine.c24,
               'c25': routine.c25, 'c26': routine.c26, 'c27': routine.c27, 'c31': routine.c31, 'c32': routine.c32,
               'c33': routine.c33, 'c34': routine.c34, 'c35': routine.c35, 'c36': routine.c36, 'c37': routine.c37,
               'c41': routine.c41, 'c42': routine.c42, 'c43': routine.c43, 'c44': routine.c44, 'c45': routine.c45,
               'c46': routine.c46, 'c47': routine.c47, 'c51': routine.c51, 'c52': routine.c52, 'c53': routine.c53,
               'c54': routine.c54, 'c55': routine.c55, 'c56': routine.c56, 'c57': routine.c57, 'c61': routine.c61,
               'c62': routine.c62, 'c63': routine.c63, 'c64': routine.c64, 'c65': routine.c65, 'c66': routine.c66,
               'c67': routine.c67, 'c71': routine.c71, 'c72': routine.c72, 'c73': routine.c73, 'c74': routine.c74,
               'c75': routine.c75, 'c76': routine.c76, 'c77': routine.c77, 'c81': routine.c81, 'c82': routine.c82,
               'c83': routine.c83, 'c84': routine.c84, 'c85': routine.c85, 'c86': routine.c86, 'c87': routine.c87,
               'c91': routine.c91, 'c92': routine.c92, 'c93': routine.c93, 'c94': routine.c94, 'c95': routine.c95,
               'c96': routine.c96, 'c97': routine.c97, 'c101': routine.c101, 'c102': routine.c102, 'c103': routine.c103,
               'c104': routine.c104, 'c105': routine.c105, 'c106': routine.c106, 'c107': routine.c107, 'schoolName': routine.schoolName}
    return render(request, 'main/routinePrint.html', context)
